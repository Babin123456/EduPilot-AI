"""Personal file storage routes — upload, list, download, delete."""

from __future__ import annotations

import base64
import io
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, Response
from fastapi.responses import FileResponse, StreamingResponse
from pymongo.database import Database

from app.api.deps import get_current_teacher
from app.core.config import get_settings
from app.core.database import get_db

router = APIRouter()

ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".csv", ".xlsx", ".xls",
    ".ppt", ".pptx", ".txt", ".png", ".jpg", ".jpeg", ".webp", ".gif"
}
MAX_FILE_SIZE = 15 * 1024 * 1024  # 15 MB


def _personal_files_dir(teacher_id: str) -> Path:
    """Return (and create) the personal-files directory for a teacher."""
    settings = get_settings()
    base = Path(settings.storage_local_path).resolve() / "personal_files" / teacher_id
    base.mkdir(parents=True, exist_ok=True)
    return base


@router.post("/upload")
async def upload_personal_file(
    file: UploadFile = File(...),
    teacher: dict = Depends(get_current_teacher),
    db: Database = Depends(get_db),
):
    """Upload a personal file to the teacher's private vault."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided.")

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: PDF, DOCX, CSV, XLSX, PPT, TXT, Images.",
        )

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File must be smaller than 15 MB.")

    file_id = uuid.uuid4().hex
    safe_filename = f"{file_id}{ext}"

    # Try saving to disk
    try:
        dest = _personal_files_dir(teacher["id"]) / safe_filename
        dest.write_bytes(content)
    except Exception as exc:
        print(f"[PersonalFiles] Disk write warning: {exc}")

    # Base64 encode for database fallback (ensures availability in serverless environments)
    file_b64 = base64.b64encode(content).decode("utf-8")

    # Extract text content for doc/pdf/txt viewing without downloading
    extracted_text = ""
    file_type_str = ext.lstrip(".").lower()
    if file_type_str in ("docx", "doc"):
        try:
            import docx
            doc_file = docx.Document(io.BytesIO(content))
            paras = [p.text for p in doc_file.paragraphs if p.text.strip()]
            extracted_text = "\n\n".join(paras).strip()
        except Exception as e:
            print(f"[PersonalFiles] DOCX text extraction error: {e}")
    elif file_type_str == "pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = [p.extract_text() or "" for p in reader.pages[:15]]
            extracted_text = "\n\n".join(pages).strip()
        except Exception as e:
            print(f"[PersonalFiles] PDF text extraction error: {e}")
    elif file_type_str in ("txt", "csv"):
        try:
            extracted_text = content.decode("utf-8", errors="ignore").strip()
        except Exception:
            pass

    settings = get_settings()
    download_url = f"{settings.backend_url}/api/v1/personal-files/download/{file_id}"

    # Store metadata in MongoDB
    doc = {
        "id": file_id,
        "teacher_id": teacher["id"],
        "original_filename": file.filename,
        "stored_filename": safe_filename,
        "file_type": file_type_str,
        "file_size_bytes": len(content),
        "download_url": download_url,
        "file_b64": file_b64,
        "extracted_text": extracted_text,
        "created_at": datetime.now(timezone.utc),
    }
    db.teacher_personal_files.insert_one(doc)

    # Auto-ingest profile file into RAG Knowledge Base if PDF or DOCX
    if ext in (".pdf", ".docx"):
        from app.services.rag_service import ingest_document
        try:
            ingest_document(
                file_bytes=content,
                filename=file.filename,
                teacher_id=teacher["id"],
                db=db,
            )
        except Exception as e:
            print(f"[Warning] RAG auto-ingest failed for personal file: {e}")

    return {
        "id": doc["id"],
        "original_filename": doc["original_filename"],
        "file_type": doc["file_type"],
        "file_size_bytes": doc["file_size_bytes"],
        "download_url": doc["download_url"],
        "created_at": doc["created_at"].isoformat(),
    }


@router.get("/")
def list_personal_files(
    teacher: dict = Depends(get_current_teacher),
    db: Database = Depends(get_db),
):
    """List all personal files for the authenticated teacher."""
    files = list(
        db.teacher_personal_files.find(
            {"teacher_id": teacher["id"]},
            {"file_b64": 0, "_id": 0},
        ).sort("created_at", -1).limit(100)
    )
    for f in files:
        if isinstance(f.get("created_at"), datetime):
            f["created_at"] = f["created_at"].isoformat()
    return files


@router.get("/download/{file_id}")
def download_personal_file(
    file_id: str,
    teacher: dict = Depends(get_current_teacher),
    db: Database = Depends(get_db),
):
    """Download a specific personal file safely."""
    doc = db.teacher_personal_files.find_one({"id": file_id})
    if not doc:
        raise HTTPException(status_code=404, detail="File not found.")

    if doc.get("teacher_id") != teacher["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this file.")

    filename = doc.get("original_filename", "downloaded_file")
    # Sanitize header filename
    safe_filename = filename.replace('"', '\\"').replace("\n", "").replace("\r", "")

    # 1. Serve from MongoDB Base64 if available
    if doc.get("file_b64"):
        try:
            b64_str = doc["file_b64"]
            missing_padding = len(b64_str) % 4
            if missing_padding:
                b64_str += "=" * (4 - missing_padding)
            raw_bytes = base64.b64decode(b64_str)
            return Response(
                content=raw_bytes,
                media_type="application/octet-stream",
                headers={"Content-Disposition": f'attachment; filename="{safe_filename}"'},
            )
        except Exception as err:
            print(f"[PersonalFiles] Base64 decode warning for {file_id}: {err}")

    # 2. Serve from local disk fallback
    try:
        stored_name = doc.get("stored_filename", f"{file_id}.pdf")
        file_path = _personal_files_dir(teacher["id"]) / stored_name
        if file_path.exists():
            return FileResponse(
                path=str(file_path),
                filename=safe_filename,
                media_type="application/octet-stream",
            )
    except Exception as err:
        print(f"[PersonalFiles] Local disk read warning for {file_id}: {err}")

    raise HTTPException(status_code=404, detail="File data is unavailable.")


@router.delete("/{file_id}")
def delete_personal_file(
    file_id: str,
    teacher: dict = Depends(get_current_teacher),
    db: Database = Depends(get_db),
):
    """Delete a personal file from storage and database."""
    doc = db.teacher_personal_files.find_one({
        "id": file_id,
        "teacher_id": teacher["id"],
    })
    if not doc:
        raise HTTPException(status_code=404, detail="File not found.")

    # Remove from disk if exists
    try:
        file_path = _personal_files_dir(teacher["id"]) / doc["stored_filename"]
        if file_path.exists():
            os.unlink(file_path)
    except Exception:
        pass

    # Remove from DB
    db.teacher_personal_files.delete_one({"id": file_id})
    return {"message": "File deleted successfully."}


MIME_MAP = {
    "pdf": "application/pdf",
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "webp": "image/webp",
    "gif": "image/gif",
    "txt": "text/plain; charset=utf-8",
    "csv": "text/plain; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "doc": "application/msword",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "ppt": "application/vnd.ms-powerpoint",
}


@router.get("/view/{file_id}")
def view_personal_file_inline(
    file_id: str,
    db: Database = Depends(get_db),
):
    """Serve personal file inline for native browser rendering in a new window/tab."""
    doc = db.teacher_personal_files.find_one({"id": file_id})
    if not doc:
        raise HTTPException(status_code=404, detail="File not found.")

    file_type = doc.get("file_type", "").lower()
    mime_type = MIME_MAP.get(file_type, "application/pdf" if file_type == "pdf" else "application/octet-stream")
    filename = doc.get("original_filename", "file")
    safe_filename = filename.replace('"', '\\"').replace("\n", "").replace("\r", "")

    # 1. Serve from Base64
    if doc.get("file_b64"):
        try:
            b64_str = doc["file_b64"]
            missing_padding = len(b64_str) % 4
            if missing_padding:
                b64_str += "=" * (4 - missing_padding)
            raw_bytes = base64.b64decode(b64_str)
            return Response(
                content=raw_bytes,
                media_type=mime_type,
                headers={"Content-Disposition": f'inline; filename="{safe_filename}"'},
            )
        except Exception as err:
            print(f"[PersonalFiles] Base64 inline decode warning: {err}")

    # 2. Local disk fallback
    try:
        teacher_id = doc.get("teacher_id", "")
        stored_name = doc.get("stored_filename", f"{file_id}.{file_type}")
        file_path = _personal_files_dir(teacher_id) / stored_name
        if file_path.exists():
            return FileResponse(
                path=str(file_path),
                filename=safe_filename,
                media_type=mime_type,
                headers={"Content-Disposition": f'inline; filename="{safe_filename}"'},
            )
    except Exception as err:
        print(f"[PersonalFiles] Disk inline read warning: {err}")

    raise HTTPException(status_code=404, detail="File content is unavailable.")


@router.get("/text-content/{file_id}")
def get_personal_file_text_content(
    file_id: str,
    teacher: dict = Depends(get_current_teacher),
    db: Database = Depends(get_db),
):
    """Retrieve extracted text content from document (DOC, DOCX, PDF, TXT) for inline viewing without downloading."""
    doc = db.teacher_personal_files.find_one({"id": file_id})
    if not doc:
        raise HTTPException(status_code=404, detail="File not found.")

    if doc.get("teacher_id") != teacher["id"]:
        raise HTTPException(status_code=403, detail="Not authorized.")

    extracted_text = doc.get("extracted_text", "")

    # Extract text on-the-fly if missing
    if not extracted_text:
        raw_bytes = None
        if doc.get("file_b64"):
            try:
                b64_str = doc["file_b64"]
                missing_padding = len(b64_str) % 4
                if missing_padding:
                    b64_str += "=" * (4 - missing_padding)
                raw_bytes = base64.b64decode(b64_str)
            except Exception:
                pass

        file_type = doc.get("file_type", "").lower()
        if raw_bytes:
            if file_type in ("docx", "doc"):
                try:
                    import docx
                    doc_file = docx.Document(io.BytesIO(raw_bytes))
                    paras = [p.text for p in doc_file.paragraphs if p.text.strip()]
                    extracted_text = "\n\n".join(paras).strip()
                except Exception as e:
                    print(f"DOCX extract error: {e}")
            elif file_type == "pdf":
                try:
                    import pypdf
                    reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
                    pages = [p.extract_text() or "" for p in reader.pages[:15]]
                    extracted_text = "\n\n".join(pages).strip()
                except Exception as e:
                    print(f"PDF extract error: {e}")
            elif file_type in ("txt", "csv"):
                try:
                    extracted_text = raw_bytes.decode("utf-8", errors="ignore").strip()
                except Exception:
                    pass

        if extracted_text:
            db.teacher_personal_files.update_one({"id": file_id}, {"$set": {"extracted_text": extracted_text}})

    return {
        "id": doc["id"],
        "original_filename": doc.get("original_filename", ""),
        "file_type": doc.get("file_type", ""),
        "extracted_text": extracted_text or f"Document Summary: {doc.get('original_filename')} is stored in your private vault.",
    }
